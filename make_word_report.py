# make_word_report.py
import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = os.path.abspath(os.path.dirname(__file__))
DELIVERABLE_DIR = os.path.join(ROOT, "deliverable")
MANIFEST_PATH = os.path.join(DELIVERABLE_DIR, "manifest.json")
CAPTIONS_PATH = os.path.join(DELIVERABLE_DIR, "captions.md")
OUTPUT_DOC = os.path.join(ROOT, "AI_Trial_Submission.docx")

CODE_FILES = ["app.py", "model.py", "utils.py"]
README_PATH = os.path.join(ROOT, "README.md")

def load_manifest(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def read_text(path):
    if not os.path.exists(path): return None
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def insert_image_with_caption(doc, img_path, title, caption, meta):
    if not os.path.exists(img_path):
        add_paragraph(doc, f"[Missing image file: {img_path}]", bold=True)
        return
    # Insert title
    add_paragraph(doc, title, bold=True, size=12)
    # Insert image (fit to ~5.5 inches width)
    try:
        doc.add_picture(img_path, width=Inches(5.5))
    except Exception as e:
        add_paragraph(doc, f"[Could not insert image {img_path} — {e}]", bold=True)
    # Caption and metadata
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("Caption: ")
    run.bold = True
    p.add_run(caption + "\n")
    run2 = p.add_run("Metadata: ")
    run2.bold = True
    meta_text = json.dumps(meta, indent=2, ensure_ascii=False)
    p.add_run(meta_text)
    doc.add_paragraph()  # blank line

def embed_code_files(doc, files):
    add_heading(doc, "Included code files", level=2)
    for f in files:
        path = os.path.join(ROOT, f)
        if not os.path.exists(path):
            add_paragraph(doc, f"- {f} (not found)", bold=False)
            continue
        add_paragraph(doc, f"File: {f}", bold=True)
        code_text = read_text(path)
        # Insert as preformatted text block
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()  # spacer

def build_document():
    if not os.path.exists(MANIFEST_PATH):
        print("manifest.json not found in deliverable/. Create it and re-run.")
        return

    manifest = load_manifest(MANIFEST_PATH)
    captions_text = read_text(CAPTIONS_PATH) or ""
    readme_text = read_text(README_PATH) or ""

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    add_heading(doc, "AI Trial — Submission Document", level=1)
    add_paragraph(doc, "Name: Sneha")
    add_paragraph(doc, "Email: singhsneha115506@gmail.com")
    add_paragraph(doc, "Brief: This document contains code excerpts, generation metadata, and the sample images produced for the 3-article trial task.")
    doc.add_paragraph("")

    # README section
    if readme_text:
        add_heading(doc, "README (summary)", level=2)
        # Add first ~500 words of README
        snippet = "\n".join(readme_text.splitlines()[:80])
        add_paragraph(doc, snippet)

    # Images
    add_heading(doc, "Generated Images and Captions", level=2)
    for entry in manifest:
        file_rel = entry.get("file")
        img_path = os.path.join(DELIVERABLE_DIR, file_rel)
        title = entry.get("title", file_rel)
        # find caption from captions.md (simple search)
        caption = ""
        try:
            # parse captions.md naive: look for the file path line and read next lines until blank
            if captions_text:
                lines = captions_text.splitlines()
                idx = None
                for i, line in enumerate(lines):
                    if line.strip().endswith(file_rel):
                        idx = i
                        break
                if idx is not None:
                    # read next non-empty lines as title/caption
                    # next line title, then 'Caption: ...' lines
                    # we'll assemble a short caption if present
                    # safer: search for "Caption:" after idx
                    cap = ""
                    for j in range(idx, min(idx+8, len(lines))):
                        if lines[j].strip().lower().startswith("caption:"):
                            cap = lines[j].split(":",1)[1].strip()
                            break
                    caption = cap
        except Exception:
            caption = ""

        if not caption:
            caption = entry.get("notes", "")

        insert_image_with_caption(doc, img_path, title, caption, {
            "prompt": entry.get("prompt"),
            "negative_prompt": entry.get("negative_prompt"),
            "steps": entry.get("steps"),
            "guidance": entry.get("guidance"),
            "size": entry.get("size"),
            "article": entry.get("article")
        })

    # Code files
    embed_code_files(doc, CODE_FILES)

    # Manifest appendix
    add_heading(doc, "Manifest (full metadata)", level=2)
    add_paragraph(doc, json.dumps(manifest, indent=2, ensure_ascii=False))

    # Save
    doc.save(OUTPUT_DOC)
    print("Saved Word file:", OUTPUT_DOC)

if __name__ == "__main__":
    build_document()
